import io

from docx import Document
from fastapi.testclient import TestClient

from backend.app.main import app

client=TestClient(app)


def headers():
    with client:
        response=client.post('/api/auth/login',json={'email':'recruiter@recruitai.local','password':'recruitai'})
    return {'Authorization':f"Bearer {response.json()['access_token']}"}


def resume():
    output=io.BytesIO();doc=Document();doc.add_heading('Katherine Johnson');doc.add_paragraph('katherine@example.com');doc.add_paragraph('Python analytics engineering mathematics systems');doc.save(output);output.seek(0);return output


def test_job_distribution_and_invitation_outbox():
    auth=headers()
    job=client.post('/api/jobs',headers=auth,json={'title':'Data Engineer','department':'Data','location':'Remote','description':'Build Python data systems and analytics platforms.'}).json()
    draft_post=client.post(f"/api/jobs/{job['id']}/postings",headers=auth,json={'boards':['linkedin']})
    assert draft_post.status_code==409
    client.patch(f"/api/jobs/{job['id']}/publish",headers=auth)
    confirmations=client.post(f"/api/jobs/{job['id']}/postings",headers=auth,json={'boards':['linkedin','indeed']})
    assert confirmations.status_code==200
    assert {item['board'] for item in confirmations.json()}=={'linkedin','indeed'}
    assert len(client.get(f"/api/jobs/{job['id']}/postings",headers=auth).json())==2
    application=client.post(f"/api/jobs/{job['id']}/applications",headers=auth,data={'name':'Katherine Johnson','email':'katherine@example.com'},files={'resume':('resume.docx',resume(),'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}).json()
    client.post(f"/api/applications/{application['id']}/interview",headers=auth)
    delivery=client.post(f"/api/applications/{application['id']}/interview/send-invite",headers=auth)
    assert delivery.status_code==200
    assert delivery.json()['recipient']=='katherine@example.com'
    assert delivery.json()['provider']=='local-outbox'
    communications=client.get(f"/api/applications/{application['id']}/communications",headers=auth).json()
    assert communications[0]['status']=='sent'
    assert 'Secure interview link' in communications[0]['body']
