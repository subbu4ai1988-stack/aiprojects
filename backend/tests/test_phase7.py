import io
from datetime import UTC, datetime, timedelta
from uuid import uuid4

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import select

from backend.app.database import SessionLocal
from backend.app.main import app
from backend.app.models import Candidate
from backend.app.privacy_models import CandidatePrivacy, PrivacyAuditLog

client = TestClient(app)


def login(email: str, password: str) -> dict[str, str]:
    with client:
        response = client.post("/api/auth/login", json={"email": email, "password": password})
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['access_token']}"}


def resume_file() -> io.BytesIO:
    buffer = io.BytesIO()
    document = Document()
    document.add_heading("Privacy Test Candidate")
    document.add_paragraph("privacy@example.com")
    document.add_paragraph("Python Docker PostgreSQL API operations")
    document.save(buffer)
    buffer.seek(0)
    return buffer


def test_candidate_privacy_export_hold_and_retention_deletion():
    admin = login("admin@recruitai.local", "recruitai-admin")
    recruiter = login("recruiter@recruitai.local", "recruitai")
    assert client.get("/api/admin/privacy/summary", headers=recruiter).status_code == 403

    suffix = uuid4().hex[:8]
    email = f"privacy-{suffix}@example.com"
    job = client.post(
        "/api/jobs",
        headers=recruiter,
        json={
            "title": f"Privacy Engineer {suffix}",
            "department": "Compliance",
            "location": "Remote",
            "description": "Build privacy-aware Python and PostgreSQL recruitment services.",
            "ranking_params": {"required_skills": ["Python", "PostgreSQL"]},
        },
    ).json()
    application = client.post(
        f"/api/jobs/{job['id']}/applications",
        headers=recruiter,
        data={"name": "Privacy Candidate", "email": email, "consent": "true"},
        files={"resume": ("privacy.docx", resume_file(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert application.status_code == 200

    with SessionLocal() as db:
        candidate = db.scalar(select(Candidate).where(Candidate.email == email))
        assert candidate is not None
        candidate_id = candidate.id
        resume_path = candidate.resume_path
        privacy = db.scalar(select(CandidatePrivacy).where(CandidatePrivacy.candidate_id == candidate_id))
        assert privacy is not None
        assert privacy.consent_status == "granted"

    summary = client.get("/api/admin/privacy/summary", headers=admin)
    assert summary.status_code == 200
    assert summary.json()["retention_days"] >= 1
    export = client.post(f"/api/admin/privacy/candidates/{candidate_id}/export", headers=admin)
    assert export.status_code == 200
    assert export.json()["candidate"]["email"] == email
    assert export.json()["applications"][0]["id"] == application.json()["id"]

    hold = client.put(f"/api/admin/privacy/candidates/{candidate_id}", headers=admin, json={"legal_hold": True})
    assert hold.status_code == 200
    blocked = client.request("DELETE", f"/api/admin/privacy/candidates/{candidate_id}", headers=admin, json={"reason": "privacy test deletion"})
    assert blocked.status_code == 409
    client.put(f"/api/admin/privacy/candidates/{candidate_id}", headers=admin, json={"legal_hold": False})

    with SessionLocal() as db:
        privacy = db.scalar(select(CandidatePrivacy).where(CandidatePrivacy.candidate_id == candidate_id))
        privacy.retention_expires_at = datetime.now(UTC).replace(tzinfo=None) - timedelta(days=1)
        db.commit()

    preview = client.post("/api/admin/privacy/retention/run", headers=admin, json={"dry_run": True})
    assert preview.status_code == 200
    assert candidate_id in preview.json()["candidate_ids"]
    enforced = client.post("/api/admin/privacy/retention/run", headers=admin, json={"dry_run": False})
    assert enforced.status_code == 200
    assert enforced.json()["deleted"] >= 1

    with SessionLocal() as db:
        assert db.get(Candidate, candidate_id) is None
        actions = set(db.scalars(select(PrivacyAuditLog.action).where(PrivacyAuditLog.subject_ref.is_not(None))))
        assert {"candidate_export", "privacy_update", "retention_delete"}.issubset(actions)

    from pathlib import Path
    assert not Path(resume_path).exists()
