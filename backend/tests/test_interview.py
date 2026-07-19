import io

from docx import Document
from fastapi.testclient import TestClient

from backend.app.main import app

client = TestClient(app)


def auth():
    with client:
        response = client.post('/api/auth/login', json={'email':'recruiter@recruitai.local','password':'recruitai'})
    return {'Authorization':f"Bearer {response.json()['access_token']}"}


def resume_file():
    buffer=io.BytesIO();document=Document();document.add_heading('Grace Hopper');document.add_paragraph('grace@example.com');document.add_paragraph('8 years Python FastAPI PostgreSQL scalable systems engineering');document.save(buffer);buffer.seek(0);return buffer


def test_edit_record_rerecord_and_feedback_workflow():
    headers=auth()
    job=client.post('/api/jobs',headers=headers,json={'title':'Platform Engineer','department':'Engineering','location':'Remote','description':'Build scalable Python FastAPI and PostgreSQL platform services.','ranking_params':{'required_skills':['Python','FastAPI']}}).json()
    application=client.post(f"/api/jobs/{job['id']}/applications",headers=headers,data={'name':'Grace Hopper','email':'grace@example.com'},files={'resume':('grace.docx',resume_file(),'application/vnd.openxmlformats-officedocument.wordprocessingml.document')})
    assert application.status_code==200
    application_id=application.json()['id']
    setup=client.post(f'/api/applications/{application_id}/interview',headers=headers)
    assert setup.status_code==200
    token=setup.json()['token']
    questions=[{'text':'Explain a scalable FastAPI design.','category':'Technical','difficulty':'Complex'}]
    updated=client.put(f'/api/applications/{application_id}/interview/questions',headers=headers,json={'questions':questions})
    assert updated.json()['status']=='prepared'
    for attempt in range(2):
        video=client.post(f'/api/interviews/{token}/answers/0/video',files={'video':('answer.webm',io.BytesIO(b'local-video-data'),'video/webm')})
        assert video.status_code==200
        assert video.json()['recording_attempts']==attempt+1
    blocked=client.post(f'/api/interviews/{token}/answers/0/video',files={'video':('answer.webm',io.BytesIO(b'third'),'video/webm')})
    assert blocked.status_code==409
    submitted=client.post(f'/api/interviews/{token}/answers',json={'answers':[{'question':questions[0]['text'],'answer':'I use async FastAPI services, PostgreSQL, queues, metrics, and horizontal scaling.'}]})
    assert submitted.status_code==200
    feedback=client.get(f'/api/applications/{application_id}/feedback',headers=headers)
    assert feedback.status_code==200
    assert feedback.json()['recommendation'] in {'select','consider','reject'}
    assert feedback.json()['answers'][0]['video_url'].startswith('/api/storage/download?')
